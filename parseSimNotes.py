import os, sys
import docx
import re
from glob import glob
import datetime

# Hardcoded settings
simNotesDir = "Y:\\hsud3\\sim_notes"
patientFileName = "G:\\srspre\\newCases_2020-05-01.txt"
# Lower case strings to search for in the CT sim site, if it matches any of these it counts.
brainSiteIdentifiers = ['brain','gyrus','srs','wbrt','cerebellum','occipital','frontal','lobe','temp','cere','lesion','ganglia','mets','brian','thalamus','plexus','hippocampus','cerebellar','parietal']
gracePeriod = 7 # Days around (Ct sim date, Txt start date) to search for the right sim note
hardcodedContrastTypes = ['None','Rectal','IV Contrast','Vaginal','Oral','Full Catheter']
outputFileDelimiter = ';'

gracePeriodDt = datetime.timedelta(days=gracePeriod)
patientFile = open(patientFileName,"r")
nCase=0
nPatientsWithNote=0
nPatientsWithDatedNote=0
patientsWithoutNote=[]
outputFileName = "parsedSimNotes_"+os.path.basename(patientFileName)
outputFile = open(outputFileName,"w+")
headerRow = ['MRN','Hash','mrDate','ctSimDate','txtStartDate','simNoteDate','simNoteSite']
for contrastType in hardcodedContrastTypes:
  headerRow.append('contrast'+contrastType.replace(' ',''))
outputFile.write(outputFileDelimiter.join(headerRow)+"\n")

for line in patientFile:
  
  # Get patient information from the text file and assign it to patientFileTokens
  patientFileTokens = line.split()
  assert len(patientFileTokens) == 5, "Rows must be: MRN Hash, txtStartDate mrDate ctSimDate"
  mrn = patientFileTokens[0]
  hash = patientFileTokens[1]
  txtStartDate = patientFileTokens[2]
  mrDate       = patientFileTokens[3]
  ctSimDate    = patientFileTokens[4]
  txtStartDateTokens = [int(i) for i in txtStartDate.split('/')]
  txtStartDateObj = datetime.date(txtStartDateTokens[2], txtStartDateTokens[0], txtStartDateTokens[1])
  mrDateTokens = [int(i) for i in mrDate.split('/')]
  mrDateObj = datetime.date(mrDateTokens[2], mrDateTokens[0], mrDateTokens[1])
  ctSimDateTokens = [int(i) for i in ctSimDate.split('/')]
  ctSimDateObj = datetime.date(ctSimDateTokens[2], ctSimDateTokens[0], ctSimDateTokens[1])
  windowStart = ctSimDateObj - gracePeriodDt
  windowEnd   = txtStartDateObj + gracePeriodDt
 
  nCase+=1
  patientSimNoteDirs = glob(os.path.join(simNotesDir, mrn))
  print('Found', len(patientSimNoteDirs), 'dirs for mrn', mrn)
  assert(len(patientSimNoteDirs)<=1)
  if len(patientSimNoteDirs) == 0:
  #  print('  skipping')
  #  continue
    patientSimNotes = []
  else:
    patientSimNotes = set(glob(os.path.join(patientSimNoteDirs[0], "*.docx"))) - set(glob(os.path.join(patientSimNoteDirs[0], "~*.docx")))
  
  print('  Found', len(patientSimNotes), 'notes, trying to match to date', ctSimDate)
  
  foundGoodSimNote = False
  goodSimNoteSite = False
  goodSimNoteDate = False
  nMatchingSimNotes = 0
  for n, note in enumerate(patientSimNotes):
    print('   Note #%d: "%s"'%(n+1,note))
    doc = docx.Document(note)
    #for paragraph in doc.paragraphs:
      #if 'IV Contrast' in paragraph.text:
      #print(paragraph.text)
      #for run in paragraph.runs:
      #  print(run._r.xml)
    
    docSimDate = False
    # sample output from doc.tables:
    '''MEMORIAL HOSPITAL FOR CANCER AND ALLIED DISEASES
    
    
    Department of Radiation Oncology
    Simulation Note
    (with RTT Time-Out)
    99999999
    MRN
    HSU, DYLAN
    
    MEMORIAL HOSPITAL FOR CANCER AND ALLIED DISEASES
    
    
    Department of Radiation Oncology
    Simulation Note
    (with RTT Time-Out)
    DOB: 3/3/1993
    SEX:  Male
    Primary MD:  Joe Schoe
    Approving MD:  Jane Doe, MD
    DOB: 3/3/1993
    SEX:  Male
    Primary MD:  Joe Schmoe
    Approving MD:  Jane Doe, MD
    SIMULATION DATE:  5/21/2020'''
    isNotVerification = False
    for table in doc.tables:
     for row in table.rows:
      for cell in row.cells:
       for para in cell.paragraphs:
        if 'SIMULATION DATE' in para.text:
         
         words = para.text.strip().split(' ')
         docSimDate = words[len(words)-1]
         print('    Sim date:',docSimDate)
         if '/' not in docSimDate:
           docSimDate='1/1/1900'
         docSimDateTokens = [int(i) for i in docSimDate.split('/')]
         docSimDateObj = datetime.date(docSimDateTokens[2], docSimDateTokens[0], docSimDateTokens[1])
         isNotVerification = True
         break

    if isNotVerification is False:
      print('    Could not find text "SIMULATION DATE", assume verification sim note. Skipping...')
      continue
    if docSimDateObj > windowEnd or docSimDateObj < windowStart:
      print('    Sim note date is outside the range (', windowStart,',',windowEnd,'). Skipping...')
      continue

    # Now check sim Site
    isBrain = False
    docSite = False
    for para in doc.paragraphs:
     if 'Confirm Site' in para.text:
      docSite = re.sub('.*Confirm Site\s+','',para.text)
      break
     if 'Disease Site' in para.text:
      docSite = re.sub('.*Disease Site:\s+','',para.text)
      break
    docSite = docSite.strip()
    assert docSite is not False, "Couldn't find a sim site."
    print('    Site: ',docSite)
    
    # Below block is commented out because the site naming is quite inconsistent.
    # e.g. neither "right post central gyrus" nor "lt frontal" will match for "brain"
    
    for identifier in brainSiteIdentifiers:
      if identifier in docSite.lower():
        isBrain = True
        break
    if docSite == '':
      isBrain = True
    if not isBrain:
      print('    Not a brain sim note, skipping...')
      continue
      
    # Find the paragraphs that have contrast information.
    # Newer sim notes tend to begin the following section with 'Confirm Special Procedures'
    # Older sim notes have a few checkbox sections, then 'Skin Marking'.
    # For the older notes, we can't see the 'None' box twice, so stop after the first sections.
    contrastSectionBegin = False
    contrastSectionEnd = False
    for p, para in enumerate(doc.paragraphs):
      if contrastSectionBegin is False and ('Confirm Contrast' in para.text or 'Contrast:' in para.text):
        contrastSectionBegin = p
      if contrastSectionEnd is False and ('Confirm Special Procedures' in para.text or 'Additional Contrast' in para.text):
        contrastSectionEnd = p
    assert contrastSectionBegin and contrastSectionEnd, "Malformed contrast section, fix this"
    
    # Now get the contrast checkboxes.
    contrastTypeCheckboxes={}
    
    # Note: Older sim notes can have different contrast checkboxes
    for contrastType in hardcodedContrastTypes:
     foundContrastTypeParagraph = False
     contrastTypeCheckboxes[contrastType] = False
     for para in doc.paragraphs[contrastSectionBegin:contrastSectionEnd]:
      if contrastType in para.text:
       foundContrastTypeParagraph = True
       contrastXmlBegin = para._p.xml.find(contrastType)
       xmlFindCheck = para._p.xml[:contrastXmlBegin]
       xmlFindCheckLines = xmlFindCheck.splitlines()
       
       # We're looking for something like this:
       '''
       <w:fldChar w:fldCharType="begin">
         <w:ffData>
           <w:name w:val="Check1"/>
           <w:enabled/>
           <w:calcOnExit w:val="0"/>
           <w:checkBox>
             <w:sizeAuto/>
             <w:default w:val="0"/>
             <w:checked/>
           </w:checkBox>
         </w:ffData>
       </w:fldChar>'''
       complexFieldCharacterBegin = -1
       complexFieldCharacterEnd = -1
       checkBoxBegin = -1
       checkBoxEnd = -1
       checkBoxCheckedTag = -1
       
       for i, line in enumerate(xmlFindCheckLines):
        if '<w:fldChar w:fldCharType="begin">' in line:
         complexFieldCharacterBegin = i
        elif '</w:fldChar>' in line:
         complexFieldCharacterEnd = i
        elif '<w:checkBox>' in line:
         checkBoxBegin = i
        elif '</w:checkBox>' in line:
         checkBoxEnd = i
        elif '<w:checked/>' in line:
         checkBoxCheckedTag = i
       assert (
         complexFieldCharacterBegin > 0 and
         checkBoxBegin > complexFieldCharacterBegin and
         checkBoxEnd > checkBoxBegin and
         complexFieldCharacterEnd > checkBoxEnd
       ), "Checkbox is malformed, fix this"
       if checkBoxCheckedTag > checkBoxBegin and checkBoxCheckedTag < checkBoxEnd:
         contrastTypeCheckboxes[contrastType] = True

       #print(xmlFindCheck)
      #print(para.text)
      #print(para._p.xml)
    print('    Contrast checkboxes:', contrastTypeCheckboxes)
    totalTrueCheckboxes=0
    for contrastType,value in contrastTypeCheckboxes.items():
      totalTrueCheckboxes+=int(value)
    assert (totalTrueCheckboxes==1 or contrastTypeCheckboxes['None']==False), "Illogical contrast checkboxes, fix the parsing code."
    if totalTrueCheckboxes==0:
      contrastTypeCheckboxes['None']=True
    
  #input("  PRESS ENTER TO CONTINUE...")
    #if ctSimDate==docSimDate:

    foundGoodSimNote = True
    goodSimNoteDate = docSimDate
    nMatchingSimNotes += 1
    goodSimNoteSite = docSite
    goodSimNoteContrastTypeCheckboxes = contrastTypeCheckboxes
    print('   Matched to Note #%d.'%(n+1))
    break
  #assert nMatchingSimNotes <= 1, "Too many matching sim notes. Review the output above and fix the search criteria."
  dataRow = patientFileTokens
  dataRow += [str(goodSimNoteDate),str(goodSimNoteSite)]
  if foundGoodSimNote:
    nPatientsWithNote += 1
    for contrastType in hardcodedContrastTypes:
      dataRow.append( str(int(contrastTypeCheckboxes[contrastType])))
  else:
    patientsWithoutNote.append(mrn)
    for contrastType in hardcodedContrastTypes:
      dataRow.append("-1")
  # sanitize output from delimiters
  for i in range(len(dataRow)):
    dataRow[i] = dataRow[i].replace(outputFileDelimiter,'')
  
  outputLine = outputFileDelimiter.join(dataRow)+"\n"
  outputFile.write(outputLine)
  
outputFile.close()
print('%d/%d had notes'%(nPatientsWithNote,nCase))
print('Patients without notes:', patientsWithoutNote)
